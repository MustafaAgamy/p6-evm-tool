"""The v5 Narrative Report writes a valid, editable .docx with real content."""
import os

from docx import Document

from p6_evm.parser import parse_file
from p6_narrative.docx_writer import write_docx
from p6_narrative.report import build_report
from tests import intel_fixtures as F

FIX = os.path.join(os.path.dirname(__file__), 'fixtures', 'minimal.xml')


def test_writes_openable_docx_from_minimal(tmp_path):
    doc = build_report(parse_file(FIX)).to_dict()
    out = os.path.join(str(tmp_path), 'narrative.docx')
    write_docx(doc, out)
    assert os.path.exists(out) and os.path.getsize(out) > 0
    text = '\n'.join(p.text for p in Document(out).paragraphs)
    assert 'Test Project' in text
    assert 'Project Overview' in text


def test_writes_editable_tables_on_a_rich_schedule(tmp_path):
    # a matrix-EPC schedule has real scope worlds -> the breakdown renders as a Word table
    doc = build_report(F.matrix_epc(4)).to_dict()
    out = os.path.join(str(tmp_path), 'rich.docx')
    write_docx(doc, out)
    reopened = Document(out)
    paras = [p.text for p in reopened.paragraphs]
    assert any('Work Breakdown Structure' in t for t in paras)
    assert any('Sequence of Work' in t for t in paras)
    # native editable table(s), not flattened images
    assert len(reopened.tables) >= 1

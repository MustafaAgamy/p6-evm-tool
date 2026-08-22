"""Write a Baseline Narrative Report (dict) to a professional, natively-editable
Word (.docx) document.

Everything lands as native, editable Word content — paragraphs, tables and outlines,
never a flattened image of text — so a planner can change any word, number, row or
cell in Word and finish it on their letterhead. The document is styled to read like a
senior planning deliverable and to MATCH the HTML / PDF export (same section structure,
one font family, the same heading sizes, navy table headers, zebra rows, an A4 page
with ~2 cm margins, a page-level border on every page, a logo/title header that repeats,
and an automatic ``Page X of Y`` footer). See ``NARRATIVE_RECONCILIATION.md`` §B.

It consumes the model produced by :func:`p6_narrative.report.build_report` → ``to_dict()``:

    {meta:{project_name, project_id, mode, data_date?, logos?}, sections:[section, …]}

and renders EVERY section kind the producer can emit:

    overview · keyvals · ms_table · timeline · value · scope · table · wbs_tree ·
    codes · idanatomy · seq · interfaces · prose · costbars · cashflow · image

The v5 kinds (overview / ms_table / wbs_tree / seq / interfaces) and the recovered
content kinds (keyvals / timeline / value / scope / table / codes / idanatomy / prose /
costbars / cashflow / image) all render here as native tables, outlines and paragraphs.
"""
import base64
import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── palette (matches the on-screen / PDF renderer) ───────────────────────────
NAVY = RGBColor(0x26, 0x5F, 0x7E)
INK = RGBColor(0x1A, 0x1D, 0x21)
GREY = RGBColor(0x8A, 0x90, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_HEADER_FILL = '265F7E'   # table header background
_ZEBRA_FILL = 'F2F6F8'    # alternating body-row background
_BORDER_CLR = '265F7E'    # page + rule colour
_ARROW = ' → '       # " → " for sequence / macro-flow lines
_FONT = 'Calibri'
_BODY_PT = 10.5
_TABLE_PT = 9.5


# ── low-level Word helpers ───────────────────────────────────────────────────
def _img_bytes(data_url):
    """Decode a 'data:image/...;base64,XXXX' URL (or bare base64) to a BytesIO."""
    if not data_url:
        return None
    try:
        b64 = data_url.split(',', 1)[1] if ',' in data_url else data_url
        return io.BytesIO(base64.b64decode(b64))
    except Exception:
        return None


def _money(v):
    try:
        return '{:,.0f}'.format(float(v))
    except (TypeError, ValueError):
        return '' if v is None else str(v)


def _count(v):
    try:
        return '{:,}'.format(int(v))
    except (TypeError, ValueError):
        return '' if v is None else str(v)


def _fmt_iso(v):
    """Render an ISO date (or any date-ish value) as 'DD Mon YYYY'; else pass through."""
    if not v:
        return ''
    if isinstance(v, datetime):
        return '%d %s %d' % (v.day, v.strftime('%b'), v.year)
    try:
        dt = datetime.strptime(str(v)[:10], '%Y-%m-%d')
        return '%d %s %d' % (dt.day, dt.strftime('%b'), dt.year)
    except (ValueError, TypeError):
        return str(v)


def _set_cell_bg(cell, hex_fill):
    """Shade a table cell (used for header + zebra striping)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _add_field(paragraph, field_code):
    """Append a Word field (e.g. PAGE / NUMPAGES) to a paragraph as a live run."""
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve'); instr.text = ' %s ' % field_code
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t'); placeholder.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r = run._r
    for el in (begin, instr, sep, placeholder, end):
        r.append(el)
    return run


def _add_page_border(section, color=_BORDER_CLR):
    """Draw a thin page-level frame around EVERY page of a Word section (w:pgBorders).

    ``offsetFrom='page'`` with a small inset keeps the frame inside the margins, and
    because it lives in the section's ``sectPr`` Word draws it once per page — it can
    never overflow onto the next page the way a bordered long <div> can in HTML.
    """
    sectPr = section._sectPr
    # Remove any prior page border so re-runs stay idempotent.
    for existing in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(existing)
    borders = OxmlElement('w:pgBorders')
    borders.set(qn('w:offsetFrom'), 'page')
    borders.set(qn('w:display'), 'allPages')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '18')     # eighths of a point ≈ 2.25pt
        el.set(qn('w:space'), '24')  # points inset from the page edge (inside the margin)
        el.set(qn('w:color'), color)
        borders.append(el)
    # Schema order: pgBorders belongs right after pgMar.
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.addnext(borders)
    else:
        sectPr.append(borders)


def _add_page_number_footer(section):
    """Centered 'Page X of Y' footer, repeated on every page (Word PAGE/NUMPAGES fields)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ''
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Page ')
    _add_field(para, 'PAGE')
    para.add_run(' of ')
    _add_field(para, 'NUMPAGES')
    for run in para.runs:
        run.font.name = _FONT
        run.font.size = Pt(9)
        run.font.color.rgb = GREY


def _para_bottom_rule(paragraph, hex_color='C9D6DE', sz='6'):
    """Add a full-width bottom border to a paragraph (a light divider rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(document, headers, rows, widths=None, bold_last_row=False):
    """A styled, editable Word table: navy header + thin borders + zebra body rows."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, head in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(str(head))
        run.bold = True
        run.font.name = _FONT
        run.font.size = Pt(_TABLE_PT)
        run.font.color.rgb = WHITE
        _set_cell_bg(hdr_cells[i], _HEADER_FILL)
    last = len(rows) - 1
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            if ci >= len(cells):
                break
            run = cells[ci].paragraphs[0].add_run('' if val is None else str(val))
            run.font.name = _FONT
            run.font.size = Pt(_TABLE_PT)
            if bold_last_row and ri == last:
                run.bold = True
        if ri % 2 == 0 and not (bold_last_row and ri == last):
            for c in cells:
                _set_cell_bg(c, _ZEBRA_FILL)
    if widths:
        table.autofit = False
        for row in table.rows:
            for i, width in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = width
    return table


def _para(document, text, size=None, italic=False, bold=False, color=None, style=None):
    para = document.add_paragraph(style=style) if style else document.add_paragraph()
    run = para.add_run('' if text is None else str(text))
    run.font.name = _FONT
    run.font.size = Pt(size or _BODY_PT)
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return para


def _muted(document, text):
    return _para(document, text, italic=True, color=GREY)


def _lead(document, text):
    """A grey lead / caption line under a heading (matches the HTML '.lead')."""
    return _para(document, text, size=9.5, italic=True, color=GREY)


def _section_heading(document, number, title):
    heading = document.add_heading('', level=1)
    run = heading.add_run(('%s %s' % (number, title)).strip())
    run.bold = True
    run.font.name = _FONT
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    return heading


def _subheading(document, text):
    heading = document.add_heading('', level=2)
    run = heading.add_run(str(text) or '—')
    run.bold = True
    run.font.name = _FONT
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    return heading


# ── header / footer furniture ────────────────────────────────────────────────
def _add_logo_header(document, logos):
    """Put the three party logos into the page header (repeats on every page)."""
    header = document.sections[0].header
    header.is_linked_to_previous = False
    try:
        table = header.add_table(rows=1, cols=3, width=Inches(6.6))
    except Exception:
        return False
    placed = False
    for i, key in enumerate(('owner', 'consultant', 'contractor')):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img = _img_bytes(logos.get(key))
        if img:
            try:
                para.add_run().add_picture(img, width=Inches(1.5))
                placed = True
            except Exception:
                pass
    return placed


def _add_title_band_header(document, meta):
    """A repeating title band in the header (used when no party logos are supplied)."""
    header = document.sections[0].header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.text = ''
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = meta.get('project_name') or 'Project'
    run = para.add_run(str(name))
    run.bold = True
    run.font.name = _FONT
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    sep = para.add_run('   ·   Baseline Narrative Report')
    sep.font.name = _FONT
    sep.font.size = Pt(9)
    sep.font.color.rgb = GREY
    _para_bottom_rule(para)


# ── cover ────────────────────────────────────────────────────────────────────
def _add_cover(document, meta):
    for _ in range(4):
        document.add_paragraph()
    title = document.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run(meta.get('project_name') or 'Project')
    trun.bold = True
    trun.font.name = _FONT
    trun.font.size = Pt(28)
    trun.font.color.rgb = INK

    sub = document.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run('Baseline Schedule — Narrative Report')
    srun.font.name = _FONT
    srun.font.size = Pt(15)
    srun.font.color.rgb = NAVY

    rule = document.add_paragraph(); rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_bottom_rule(rule)

    bits = []
    if meta.get('project_id'):
        bits.append('Project ID: %s' % meta['project_id'])
    if meta.get('mode'):
        bits.append('Execution model: %s' % str(meta['mode']).replace('_', ' ').title())
    if meta.get('data_date'):
        bits.append('Baseline data date: %s' % meta['data_date'])
    if bits:
        line = document.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lrun = line.add_run('   ·   '.join(bits))
        lrun.font.name = _FONT
        lrun.font.size = Pt(10.5)
        lrun.font.color.rgb = GREY

    document.add_page_break()


# ── v5 section renderers ─────────────────────────────────────────────────────
def _render_overview(document, p):
    for para in p.get('paragraphs', []):
        _para(document, para)
    breakdown = p.get('breakdown') or []
    if breakdown:
        _subheading(document, 'Baseline composition')
        rows = [[b.get('world', ''), _count(b.get('count'))] for b in breakdown]
        total = p.get('total')
        if total is not None:
            rows.append(['Total', _count(total)])
        _add_table(document, ['Scope', 'Activities'], rows,
                   widths=(Inches(4.6), Inches(1.6)),
                   bold_last_row=total is not None)


def _render_ms_table(document, p, note):
    columns = p.get('columns') or ['Milestone', 'Date']
    rows = p.get('rows') or []
    if not rows:
        _muted(document, note or 'No finish milestones are defined in the file.')
        return
    _add_table(document, columns, rows, widths=(Inches(4.6), Inches(1.9)))


def _wbs_node(document, node, depth):
    """One editable outline paragraph per WBS node — indent grows with depth (full P6
    depth), level-1 bold, a '+N more' breadth marker italic-grey."""
    marker = {1: '▪  ', 2: '–  ', 3: '·  '}.get(depth, '·  ')
    para = document.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.26 * depth)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(marker + (node.get('name') or '—'))
    run.font.name = _FONT
    run.font.size = Pt(_BODY_PT)
    if node.get('more'):
        run.italic = True
        run.font.color.rgb = GREY
    elif depth <= 1:
        run.bold = True
    for child in node.get('children', []):
        _wbs_node(document, child, depth + 1)


def _render_wbs_tree(document, p):
    worlds = p.get('worlds') or []
    if not worlds:
        _muted(document, 'No work breakdown structure is defined in the file.')
        return
    _lead(document, 'The actual P6 breakdown, one editable indented outline per major '
                    'branch, at full hierarchy depth. Structure only — the execution '
                    'order is in the Sequence of Work section.')
    for w in worlds:
        root = w.get('root') or {}
        _subheading(document, root.get('name') or w.get('name') or '—')
        for child in root.get('children', []):
            _wbs_node(document, child, 1)


def _render_seq(document, p):
    worlds = p.get('worlds') or []
    if not worlds:
        _muted(document, 'No execution fronts were detected in the file.')
        return
    for w in worlds:
        _subheading(document, w.get('world') or '—')
        fronts = w.get('fronts') or []
        if not fronts:
            _muted(document, 'No work-package sequence was detected for this scope.')
            continue
        for f in fronts:
            title_p = document.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(1)
            trun = title_p.add_run(f.get('title') or 'Front')
            trun.bold = True
            trun.font.name = _FONT
            trun.font.size = Pt(_BODY_PT)

            seq = [str(s) for s in (f.get('sequence') or []) if s]
            if seq:
                sp = document.add_paragraph()
                sp.paragraph_format.left_indent = Inches(0.22)
                srun = sp.add_run(_ARROW.join(seq))
                srun.font.name = _FONT
                srun.font.size = Pt(_BODY_PT)

            meta_bits = []
            instances = [str(i) for i in (f.get('instances') or []) if i]
            if instances:
                shown = ', '.join(instances[:8])
                if len(instances) > 8:
                    shown += ', +%d more' % (len(instances) - 8)
                meta_bits.append('Applies to: ' + shown)
            acts = f.get('activities') or []
            if acts:
                meta_bits.append('%d activities' % len(acts))
            if meta_bits:
                ap = document.add_paragraph()
                ap.paragraph_format.left_indent = Inches(0.22)
                arun = ap.add_run('   ·   '.join(meta_bits))
                arun.italic = True
                arun.font.name = _FONT
                arun.font.size = Pt(9.5)
                arun.font.color.rgb = GREY


def _render_interfaces(document, p):
    macro = [str(m) for m in (p.get('macro') or []) if m]
    if macro:
        _subheading(document, 'Macro execution flow')
        mp = document.add_paragraph()
        mrun = mp.add_run(_ARROW.join(macro))
        mrun.bold = True
        mrun.font.name = _FONT
        mrun.font.size = Pt(_BODY_PT)
        mrun.font.color.rgb = NAVY
    for note in p.get('notes') or []:
        _para(document, note, style='List Bullet')
    edges = [e for e in (p.get('edges') or [])
             if isinstance(e, (list, tuple)) and len(e) >= 2]
    if edges:
        _subheading(document, 'Key building / front dependencies')
        for a, b in ((e[0], e[1]) for e in edges):
            _para(document, '%s%s%s' % (a, _ARROW, b), style='List Bullet')


# ── recovered content-breadth renderers ──────────────────────────────────────
def _render_prose(document, p):
    for para in p.get('paragraphs', []) or []:
        _para(document, para)
    for bullet in p.get('bullets', []) or []:
        _para(document, bullet, style='List Bullet')


def _render_keyvals(document, p):
    rows = [[r.get('k'), r.get('v')] for r in p.get('rows', []) or []]
    if not rows:
        _muted(document, 'No project brief fields are available.')
        return
    _add_table(document, ['Field', 'Value'], rows,
               widths=(Inches(2.4), Inches(4.2)))


def _render_image(document, p):
    img = _img_bytes(p.get('image'))
    if not img:
        return
    try:
        document.add_picture(img, width=Inches(6.2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        _muted(document, '[project layout image]')


def _render_timeline(document, p):
    items = p.get('items', []) or []
    if not items:
        _muted(document, 'No key dates are available in the file.')
        return
    rows = [[_fmt_iso(it.get('date')), it.get('label')] for it in items]
    _add_table(document, ['Date', 'Key date / milestone'], rows,
               widths=(Inches(1.9), Inches(4.7)))


def _cost_table(document, p, name_header):
    rows = [[r.get('name'), _money(r.get('cost')),
             '%s%%' % r.get('pct')] for r in p.get('rows', []) or []]
    if not rows:
        _muted(document, 'No cost-loading information is available in the file.')
        return
    total = p.get('total')
    if total is not None:
        rows.append(['Total', _money(total), '100%'])
    _add_table(document, [name_header, 'Cost', 'Share %'], rows,
               widths=(Inches(3.8), Inches(1.6), Inches(1.2)),
               bold_last_row=total is not None)


def _render_value(document, p):
    _cost_table(document, p, 'Branch')


def _render_costbars(document, p):
    _cost_table(document, p, 'WBS branch')


def _render_scope(document, p):
    stats = p.get('stats') or []
    if stats:
        labels = [s.get('l', '') for s in stats]
        values = [s.get('v', '') for s in stats]
        _add_table(document, labels, [values], bold_last_row=True)
    if p.get('intro'):
        _para(document, p['intro'])
    for b in p.get('blocks', []) or []:
        _subheading(document, b.get('discipline', '') or '—')
        if b.get('paragraph'):
            _para(document, b['paragraph'])
        for pkg in b.get('packages', []) or []:
            para = document.add_paragraph(style='List Bullet')
            run = para.add_run(str(pkg))
            run.bold = True
            run.font.name = _FONT
            run.font.size = Pt(_BODY_PT)


def _render_calendars(document, p):
    calendars = p.get('calendars', []) or []
    if calendars:
        _add_table(document, ['Calendar', 'Working days', 'Shift', 'Assigned'],
                   [[c.get('name'), c.get('working_days'), c.get('shift'),
                     _count(c.get('activities'))] for c in calendars])
    else:
        _muted(document, 'No calendars are assigned in the file.')
    holidays = p.get('holidays') or []
    if holidays:
        hp = document.add_paragraph()
        hp.paragraph_format.space_before = Pt(6)
        hrun = hp.add_run('Holidays & shutdowns')
        hrun.bold = True
        hrun.font.name = _FONT
        hrun.font.size = Pt(_BODY_PT)
        _add_table(document, ['When', 'Name', 'Days'],
                   [[h.get('range'), h.get('name'), h.get('days')] for h in holidays])


def _render_table(document, p, note):
    if p.get('view') == 'calendars':
        _render_calendars(document, p)
        return
    columns = p.get('columns') or ['—']
    rows = p.get('rows') or []
    if not rows:
        _muted(document, note or 'No rows are available for this table.')
        return
    _add_table(document, columns, rows)


def _render_codes(document, p):
    tables = p.get('tables', []) or []
    if not tables:
        _muted(document, 'No activity codes are defined in the file.')
        return
    for tbl in tables:
        _subheading(document, tbl.get('dimension', '') or '—')
        rows = [[r.get('code'), r.get('description')] for r in tbl.get('rows', []) or []]
        if rows:
            _add_table(document, ['Code', 'Description'], rows,
                       widths=(Inches(1.8), Inches(4.8)))
        else:
            _muted(document, 'No values defined for this dimension.')


def _render_idanatomy(document, p):
    aid = p.get('id', '') or ''
    lead = document.add_paragraph()
    lrun = lead.add_run('Example activity ID: ')
    lrun.font.name = _FONT
    lrun.font.size = Pt(_BODY_PT)
    idrun = lead.add_run(aid)
    idrun.bold = True
    idrun.font.name = 'Consolas'
    idrun.font.size = Pt(_BODY_PT)
    segs = [[s.get('label'), s.get('value')] for s in p.get('segments', []) or []]
    if segs:
        _add_table(document, ['Part', 'Value'], segs,
                   widths=(Inches(2.4), Inches(4.2)))


def _render_cashflow(document, p):
    points = p.get('points', []) or []
    if not points:
        _muted(document, 'Time-phased cost information is not available in the file.')
        return
    _lead(document, 'Cumulative planned cost (cost-loaded S-curve) sampled across the '
                    'baseline — illustrative of the plan.')
    rows = [[_fmt_iso(pt.get('date')), _money(pt.get('cumulative')),
             '%s%%' % pt.get('pct')] for pt in points]
    _add_table(document, ['Date', 'Cumulative cost', '% complete'], rows,
               widths=(Inches(1.9), Inches(2.6), Inches(1.6)))


_RENDER = {
    'overview': lambda doc, sec: _render_overview(doc, sec.get('payload') or {}),
    'ms_table': lambda doc, sec: _render_ms_table(doc, sec.get('payload') or {}, sec.get('note')),
    'wbs_tree': lambda doc, sec: _render_wbs_tree(doc, sec.get('payload') or {}),
    'seq': lambda doc, sec: _render_seq(doc, sec.get('payload') or {}),
    'interfaces': lambda doc, sec: _render_interfaces(doc, sec.get('payload') or {}),
    'prose': lambda doc, sec: _render_prose(doc, sec.get('payload') or {}),
    'keyvals': lambda doc, sec: _render_keyvals(doc, sec.get('payload') or {}),
    'image': lambda doc, sec: _render_image(doc, sec.get('payload') or {}),
    'timeline': lambda doc, sec: _render_timeline(doc, sec.get('payload') or {}),
    'value': lambda doc, sec: _render_value(doc, sec.get('payload') or {}),
    'scope': lambda doc, sec: _render_scope(doc, sec.get('payload') or {}),
    'table': lambda doc, sec: _render_table(doc, sec.get('payload') or {}, sec.get('note')),
    'codes': lambda doc, sec: _render_codes(doc, sec.get('payload') or {}),
    'idanatomy': lambda doc, sec: _render_idanatomy(doc, sec.get('payload') or {}),
    'costbars': lambda doc, sec: _render_costbars(doc, sec.get('payload') or {}),
    'cashflow': lambda doc, sec: _render_cashflow(doc, sec.get('payload') or {}),
}


def _render(document, section):
    handler = _RENDER.get(section.get('kind'))
    if handler is not None:
        handler(document, section)
        return
    # Graceful fallback for any unknown/future kind — never crash the export.
    payload = section.get('payload') or {}
    for para in payload.get('paragraphs', []) or []:
        _para(document, para)
    if section.get('note'):
        _muted(document, section['note'])


# ── document geometry / styles (shared spec §B) ──────────────────────────────
def _apply_base_styles(document):
    normal = document.styles['Normal']
    normal.font.name = _FONT
    normal.font.size = Pt(_BODY_PT)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    for level, size in (('Heading 1', 16), ('Heading 2', 13)):
        try:
            st = document.styles[level]
            st.font.name = _FONT
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.color.rgb = NAVY
        except KeyError:
            pass


def _apply_page_geometry(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)     # A4 portrait
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


# ── public entry point ───────────────────────────────────────────────────────
def write_docx(doc, output_path, chrome=None):
    """Render the narrative model (``doc``) to an editable .docx at ``output_path``.

    ``chrome`` is accepted for signature compatibility and ignored — this writer is
    pure ``python-docx`` and produces native, editable Word content only (every chart
    becomes an editable table / outline, so nothing is a flattened image of text).
    """
    meta = (doc or {}).get('meta', {}) or {}
    document = Document()

    _apply_base_styles(document)
    _apply_page_geometry(document.sections[0])

    # Header: party logos when supplied (repeat every page), else a title band.
    if not (meta.get('logos') and _add_logo_header(document, meta['logos'])):
        _add_title_band_header(document, meta)

    _add_cover(document, meta)

    for section in (doc or {}).get('sections', []):
        _section_heading(document, section.get('number', ''), section.get('title', ''))
        _render(document, section)

    # Page border + geometry + footer on EVERY section (one section here, but robust).
    for section in document.sections:
        _apply_page_geometry(section)
        _add_page_border(section)
        _add_page_number_footer(section)

    document.save(output_path)
    return output_path
